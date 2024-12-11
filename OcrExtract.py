import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
import os

# Caminho para o executável do Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

# Diretório tessdata
os.environ['TESSDATA_PREFIX'] = r'C:\\Program Files\\Tesseract-OCR\\tessdata'

def extract_text_from_pdf(pdf_path, output_docx):

    # Abrir o PDF
    pdf_document = fitz.open(pdf_path)
    
    # Criar um documento Word
    doc = Document()

    # Processar cada página
    for page_number in range(len(pdf_document)):
        # Extrair a página como imagem
        page = pdf_document[page_number]
        pix = page.get_pixmap()

        # Salvar a imagem temporariamente
        image_path = f"page_{page_number + 1}.png"
        pix.save(image_path)

        # Realizar OCR na imagem
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang='por')  # Língua do documento

        # Adicionar número da página ao documento
        doc.add_paragraph(f"Página {page_number + 1}")
        doc.add_paragraph(text)

        # Remover o arquivo de imagem temporário
        os.remove(image_path)

    # Salvar o documento Word
    doc.save(output_docx)
    print(f"Transcrição concluída. Arquivo salvo em: {output_docx}")

# Caminhos de entrada e saída
pdf_path = r"C:\Users\proco\OneDrive\Área de Trabalho\Qualificação\ArquivosOcr\1967.pdf"
output_docx = r"C:\Users\proco\OneDrive\Área de Trabalho\Qualificação\ArquivosExtraidos\transcricao_dissertacao67.docx"

# Executar o processo de OCR
extract_text_from_pdf(pdf_path, output_docx)

